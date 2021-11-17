# Importar as bibliotecas
import sys
import time
from Bio import SeqIO
from Bio import Entrez
from docx import Document
from docx.shared import RGBColor
from interface_grafica import doc_qrc
from interface_grafica import tela_principal
from PyQt5.QtWidgets import QApplication, QFileDialog, QWidget, QMessageBox


# Classe do Exon Finder
class ExonFinder(QWidget, tela_principal.Ui_Form):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setWindowTitle('Exon Finder')

        # Condifurações iniciais do Exon Finder
        self.sequencia = ''
        self.tempo_final_pesquisa = 0
        self.tempo_final_documento = 0

        # Botão "Pesquisar"
        self.botao_pesquisar.clicked.connect(self.funcao_pesquisar)
        # Botão "Salvar"
        self.botao_salvar.setDisabled(True)
        self.botao_salvar.clicked.connect(self.funcao_salvar)
        # Botão "?"
        self.botao_ajuda.clicked.connect(self.funcao_ajuda)
        # Botão "Versão 2.0.0"
        self.botao_versao.clicked.connect(self.funcao_versao)

    # Função do botão "Pesquisar"
    def funcao_pesquisar(self):
        inicio_exon = []
        final_exon = []
        Entrez.email = 'teste@exemplo.com'
        id_seq = self.linha_id_seq.text().upper().strip()
        with Entrez.efetch(db='nucleotide', rettype='gb', retmode='text', id=id_seq) as identificador:
            tempo_inicial_pesquisa = time.time()
            registro = SeqIO.read(identificador, 'gb')
            quantidade_features = len(registro.features)
            for item in range(quantidade_features):
                if registro.features[item].type == 'exon':
                    localizacao = list(registro.features[item].location)
                    inicio_exon.append(localizacao[0])
                    final_exon.append(localizacao[-1]+1)
            self.sequencia = str(registro.seq.lower())
            for posicao, _ in enumerate(inicio_exon):
                posicao_inicial = inicio_exon[posicao]
                posicao_final = final_exon[posicao]
                sequencia_antiga = self.sequencia[posicao_inicial:posicao_final]
                sequencia_nova = sequencia_antiga.upper()
                self.sequencia = self.sequencia.replace(sequencia_antiga, sequencia_nova)

        if id_seq != '':
            self.botao_salvar.setDisabled(False)
        self.tempo_final_pesquisa = time.time() - tempo_inicial_pesquisa

    # Função do botão "Salvar"
    def funcao_salvar(self):
        caminho = QFileDialog.getSaveFileName()[0]
        self.linha_caminho.setText(f'{caminho}.docx')
        documento = Document()
        tempo_inicial_documento = time.time()
        p = documento.add_paragraph()
        cor = RGBColor(255, 0, 0)
        for nucleotideo in self.sequencia:
            if nucleotideo.isupper():
                marcado = p.add_run(nucleotideo)
                marcado.bold = True
                marcado.font.color.rgb = cor
            else:
                p.add_run(nucleotideo)
        documento.save(f'{caminho}.docx')
        self.tempo_final_documento = time.time() - tempo_inicial_documento
        QMessageBox.information(
            self,
            'Concluído!',
            f'''Tempo de espera para conectar ao NCBI: {self.tempo_final_pesquisa:.1f} segundos.
Tempo de espera para criar o documento: {self.tempo_final_documento:.1f} segundos.
Tempo de espera total da análise: {self.tempo_final_pesquisa + self.tempo_final_documento:.1f} segundos.'''
        )
        self.botao_salvar.setDisabled(True)
        self.linha_id_seq.setText('')
        self.linha_caminho.setText('')

    # Função do botão "?"
    def funcao_ajuda(self):
        QMessageBox.information(
            self,
            'Menu de ajuda',
            '''                                         Bem vindo ao Exon Finder!
            
ID sequência: Informar o número de identificação da sequência no NCBI.

Salvar em: Escolher o diretório de destino do arquivo gerado.
Clicando no botão ao lado, será aberta uma janela para que você possa selecionar o diretório.

Para mais informações, visite: link do reposiório do Exon Finder!!!'''
        )

    # Função do botão "Versão 2.0.0"
    def funcao_versao(self):
        QMessageBox.information(
            self,
            'Dados da versão',
            '''Nome: Exon Finder
Versão: 2.0.0
Sistema Operacional: Windows
Linguagem de programação: Python 3.8
Data de lançamento: XX/XX/XX
Chave da versão: XXXXXXXXXXXXX
Desenvolvido por: Guilherme Trevisan Linhares
GitHub: link do repositório do Exon Finder!!!'''
        )


if __name__ == '__main__':
    app = QApplication(sys.argv)
    tela = ExonFinder()
    tela.show()
    sys.exit(app.exec_())
